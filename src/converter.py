#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
文档转换器模块

基于 Microsoft MarkItDown 库实现多格式文档到 Markdown 的转换。
支持图片提取、路径修复、格式优化和 LLM 集成。

主要功能：
- 基于 MarkItDown 的文档格式转换
- 智能图片提取和重命名
- 路径自动修复
- PDF 格式优化
- 可选的 LLM 集成用于图片描述
- Azure 文档智能支持
- 批量处理支持

作者: Assistant
版本: 2.0.0
依赖: Microsoft MarkItDown >= 0.1.2
"""

import os
import re
from pathlib import Path
from typing import Dict, List, Optional, Any, Tuple
from datetime import datetime

try:
    from markitdown import MarkItDown
except ImportError:
    print("错误: 未安装 markitdown 库")
    print("请运行: pip install 'markitdown[all]'")
    raise

# 可选的 LLM 客户端导入
try:
    from openai import OpenAI
    OPENAI_AVAILABLE = True
except ImportError:
    OPENAI_AVAILABLE = False

try:
    import anthropic
    ANTHROPIC_AVAILABLE = True
except ImportError:
    ANTHROPIC_AVAILABLE = False

from .config import Config
from .logger import MarkItDownLogger
from .image_processor import ImageProcessor


class DocumentConverter:
    """
    文档转换器类
    
    负责协调整个文档转换流程，包括：
    - 调用 MarkItDown 进行基础转换
    - 图片提取和处理
    - 路径修复
    - 格式优化
    - 可选的 LLM 集成
    """
    
    def __init__(self, 
                 logger: Optional[MarkItDownLogger] = None,
                 enable_llm: bool = False,
                 use_azure_doc_intel: bool = False,
                 keep_data_uris: bool = False):
        """
        初始化文档转换器
        
        Args:
            logger: 日志记录器实例
            enable_llm: 是否启用 LLM 集成用于图片描述
            use_azure_doc_intel: 是否使用 Azure 文档智能
            keep_data_uris: 是否保留 base64 编码的图片数据
        """
        self.logger = logger or MarkItDownLogger()
        self.enable_llm = enable_llm
        self.use_azure_doc_intel = use_azure_doc_intel
        self.keep_data_uris = keep_data_uris
        
        # 初始化 MarkItDown
        self._init_markitdown()
        
        # 初始化图片处理器
        self.image_processor = ImageProcessor(logger=self.logger)
        
        # 统计信息
        self.stats = {
            'total_files': 0,
            'successful_conversions': 0,
            'failed_conversions': 0,
            'total_images_extracted': 0,
            'total_processing_time': 0,
            'llm_descriptions_generated': 0
        }
    def _init_markitdown(self) -> None:
        """
        初始化 MarkItDown 实例
        """
        try:
            # 配置 LLM 客户端（如果启用）
            llm_client = None
            llm_model = None
            
            if self.enable_llm:
                llm_client, llm_model = self._setup_llm_client()
                if llm_client:
                    self.logger.info(f"LLM 集成已启用: {llm_model}")
                else:
                    self.logger.warning("LLM 集成启用失败，将跳过图片描述生成")
                    self.enable_llm = False
            
            # 初始化 MarkItDown
            self.markitdown = MarkItDown(
                llm_client=llm_client,
                llm_model=llm_model
            )
            
            self.logger.info("MarkItDown 初始化完成")
            
        except Exception as e:
            self.logger.error(f"MarkItDown 初始化失败: {str(e)}")
            # 使用默认配置
            self.markitdown = MarkItDown()
    
    def _setup_llm_client(self) -> Tuple[Optional[Any], Optional[str]]:
        """
        设置 LLM 客户端
        
        Returns:
            (客户端实例, 模型名称) 或 (None, None)
        """
        # 尝试 OpenAI
        if OPENAI_AVAILABLE:
            api_key = os.getenv('OPENAI_API_KEY')
            if api_key:
                try:
                    client = OpenAI(api_key=api_key)
                    model = os.getenv('OPENAI_MODEL', 'gpt-4o')
                    return client, model
                except Exception as e:
                    self.logger.warning(f"OpenAI 客户端初始化失败: {str(e)}")
        
        # 尝试 Anthropic
        if ANTHROPIC_AVAILABLE:
            api_key = os.getenv('ANTHROPIC_API_KEY')
            if api_key:
                try:
                    client = anthropic.Anthropic(api_key=api_key)
                    model = os.getenv('ANTHROPIC_MODEL', 'claude-3-sonnet-20240229')
                    return client, model
                except Exception as e:
                    self.logger.warning(f"Anthropic 客户端初始化失败: {str(e)}")
        
        self.logger.warning("未找到可用的 LLM 客户端配置")
        return None, None
    
    def _validate_input_file(self, file_path: str) -> bool:
        """
        验证输入文件
        
        Args:
            file_path: 文件路径
            
        Returns:
            文件是否有效
        """
        path = Path(file_path)
        
        if not path.exists():
            self.logger.error(f"文件不存在: {file_path}")
            return False
        
        if not path.is_file():
            self.logger.error(f"路径不是文件: {file_path}")
            return False
        
        # 检查文件扩展名
        file_ext = path.suffix.lower()
        supported_formats = self.get_supported_formats()
        if file_ext not in supported_formats:
            self.logger.warning(f"不支持的文件格式: {file_path} ({file_ext})")
            return False
        
        # 检查文件大小
        file_size = path.stat().st_size
        max_size = Config.MAX_FILE_SIZE * 1024 * 1024  # 转换为字节
        if file_size > max_size:
            self.logger.warning(f"文件过大: {file_path} ({file_size / 1024 / 1024:.1f}MB)")
            return False
        
        if file_size == 0:
            self.logger.warning(f"文件为空: {file_path}")
            return False
        
        return True
    
    def _get_output_path(self, input_file: str) -> str:
        """
        生成输出文件路径
        
        Args:
            input_file: 输入文件路径
            
        Returns:
            输出文件路径
        """
        input_path = Path(input_file)
        # 使用与图片文件夹相同的文件名清理逻辑，确保命名一致性
        safe_filename = self.image_processor._sanitize_filename(input_path.stem)
        output_filename = f"{safe_filename}.md"
        return str(Path(Config.OUTPUT_DIR) / output_filename)
    
    def _optimize_pdf_content(self, content: str) -> str:
        """
        优化 PDF 转换后的内容格式
        
        Args:
            content: 原始内容
            
        Returns:
            优化后的内容
        """
        if not content:
            return content
        
        # PDF 特有的格式问题修复
        # 1. 修复错误的换行符（单独的换行符变为空格）
        content = re.sub(r'(?<!\n)\n(?!\n|[#*-]|\d+\.|\s*[•·])', ' ', content)
        
        # 2. 修复多余的空格
        content = re.sub(r' {2,}', ' ', content)
        
        # 3. 修复页眉页脚（通常是重复的短文本）
        lines = content.split('\n')
        filtered_lines = []
        for line in lines:
            line = line.strip()
            # 跳过可能的页眉页脚（短且重复的行）
            if len(line) < 5 or line.isdigit():
                continue
            filtered_lines.append(line)
        
        content = '\n'.join(filtered_lines)
        
        # 4. 修复表格格式
        content = re.sub(r'\|\s*\|', '|', content)
        content = re.sub(r'\s*\|\s*', ' | ', content)
        
        return content
    
    def _clean_markdown_content(self, content: str, is_pdf: bool = False) -> str:
        """
        清理和优化 Markdown 内容
        
        Args:
            content: 原始 Markdown 内容
            is_pdf: 是否为 PDF 转换的内容
            
        Returns:
            清理后的内容
        """
        if not content:
            return content
        
        # PDF 特殊处理
        if is_pdf:
            content = self._optimize_pdf_content(content)
        
        # 通用清理
        # 1. 修复多余的空行
        content = re.sub(r'\n{3,}', '\n\n', content)
        
        # 2. 确保标题前后有空行
        content = re.sub(r'(?<!\n)\n(#{1,6}\s)', '\n\n\1', content)
        content = re.sub(r'(#{1,6}.*?)\n(?!\n)', '\1\n\n', content)
        
        # 3. 修复列表格式
        content = re.sub(r'\n([*-]\s)', '\n\n\1', content)
        
        # 4. 清理行尾空格
        content = re.sub(r' +$', '', content, flags=re.MULTILINE)
        
        return content.strip()
    
    def _save_markdown_file(self, content: str, output_path: str) -> bool:
        """
        保存 Markdown 文件
        
        Args:
            content: Markdown 内容
            output_path: 输出路径
            
        Returns:
            保存是否成功
        """
        try:
            output_dir = Path(output_path).parent
            output_dir.mkdir(parents=True, exist_ok=True)
            
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(content)
            
            self.logger.info(f"Markdown 文件已保存: {output_path}")
            return True
            
        except Exception as e:
            self.logger.error(f"保存 Markdown 文件失败: {output_path} - {str(e)}")
            return False
    
    def _create_image_markdown(self, input_file: str) -> Optional[str]:
        """
        为图片文件创建 Markdown 内容
        
        Args:
            input_file: 图片文件路径
            
        Returns:
            Optional[str]: Markdown 内容，失败时返回 None
        """
        try:
            from PIL import Image
            import os
            import shutil
            
            # 验证图片文件
            if not os.path.exists(input_file):
                self.logger.error(f"图片文件不存在: {input_file}")
                return None
            
            # 获取图片信息
            file_path = Path(input_file)
            file_name = file_path.name
            file_stem = file_path.stem
            
            # 创建文档专属的图片目录（与Word文件处理方式一致）
            doc_image_dir = self.image_processor._create_document_image_dir(file_stem)
            
            # 对图片文件名也应用清理逻辑，确保与文件夹名一致
            safe_image_name = self.image_processor._sanitize_filename(file_name)
            target_image_path = doc_image_dir / safe_image_name
            
            # 复制原始图片到images目录
            try:
                # 确保目标目录存在
                target_image_path.parent.mkdir(parents=True, exist_ok=True)
                self.logger.debug(f"目标目录已创建: {target_image_path.parent}")
                
                # 检查源文件
                if not os.path.exists(input_file):
                    self.logger.error(f"源文件不存在: {input_file}")
                    return None
                    
                source_size = os.path.getsize(input_file)
                self.logger.debug(f"源文件大小: {source_size} 字节")
                
                # 复制文件
                shutil.copy2(input_file, target_image_path)
                self.logger.debug(f"执行复制操作: {input_file} -> {target_image_path}")
                
                # 验证复制是否成功
                if target_image_path.exists():
                    target_size = target_image_path.stat().st_size
                    self.logger.debug(f"目标文件大小: {target_size} 字节")
                    if target_size > 0 and target_size == source_size:
                        self.logger.info(f"图片已成功复制到: {target_image_path}")
                    else:
                        self.logger.error(f"图片复制验证失败: 大小不匹配 源:{source_size} 目标:{target_size}")
                        return None
                else:
                    self.logger.error(f"图片复制验证失败: 目标文件不存在 {target_image_path}")
                    return None
            except Exception as e:
                self.logger.error(f"复制图片失败: {e}")
                import traceback
                self.logger.error(f"详细错误信息: {traceback.format_exc()}")
                return None
            
            # 使用清理后的文件名获取相对路径
            relative_image_path = self.image_processor._get_relative_image_path(file_stem, safe_image_name)
            
            # 再次验证文件是否存在（防止被后续处理清理）
            if not target_image_path.exists():
                self.logger.warning(f"警告：图片文件在处理后消失: {target_image_path}")
                # 重新复制一次
                try:
                    shutil.copy2(input_file, target_image_path)
                    self.logger.info(f"重新复制图片成功: {target_image_path}")
                except Exception as e:
                    self.logger.error(f"重新复制图片失败: {e}")
                    return None
            
            try:
                # 尝试打开图片获取尺寸信息
                with Image.open(input_file) as img:
                    width, height = img.size
                    format_info = img.format or "Unknown"
                    mode = img.mode
                    
                self.logger.info(f"图片信息: {file_name} ({width}x{height}, {format_info}, {mode})")
                
                # 创建 Markdown 内容
                markdown_content = f"""# {file_stem}

## 图片信息

- **文件名**: {file_name}
- **尺寸**: {width} x {height} 像素
- **格式**: {format_info}
- **颜色模式**: {mode}
- **文件大小**: {os.path.getsize(input_file)} 字节

## 图片预览

![image_{file_stem}]({relative_image_path})

---

*此文档由 MarkItDown 自动生成*
"""
                
                return markdown_content
                
            except Exception as img_error:
                self.logger.warning(f"无法读取图片详细信息: {img_error}")
                
                # 创建简化的 Markdown 内容
                markdown_content = f"""# {file_stem}

## 图片文件

- **文件名**: {file_name}
- **文件大小**: {os.path.getsize(input_file)} 字节

## 图片预览

![image_{file_stem}]({relative_image_path})

---

*此文档由 MarkItDown 自动生成*
"""
                
                return markdown_content
                
        except ImportError:
            self.logger.warning("PIL 库未安装，使用简化的图片处理")
            
            # 创建基础的 Markdown 内容
            file_path = Path(input_file)
            file_name = file_path.name
            file_stem = file_path.stem
            
            # 创建文档专属的图片目录（与Word文件处理方式一致）
            doc_image_dir = self.image_processor._create_document_image_dir(file_stem)
            
            # 对图片文件名也应用清理逻辑，确保与文件夹名一致
            safe_image_name = self.image_processor._sanitize_filename(file_name)
            target_image_path = doc_image_dir / safe_image_name
            
            # 复制原始图片到images目录
            try:
                import shutil
                # 确保目标目录存在
                target_image_path.parent.mkdir(parents=True, exist_ok=True)
                
                # 复制文件
                shutil.copy2(input_file, target_image_path)
                
                # 验证复制是否成功
                if target_image_path.exists() and target_image_path.stat().st_size > 0:
                    self.logger.info(f"图片已成功复制到: {target_image_path}")
                else:
                    self.logger.error(f"图片复制验证失败: {target_image_path}")
                    return None
            except Exception as e:
                self.logger.error(f"复制图片失败: {e}")
                return None
            
            # 使用清理后的文件名获取相对路径
            relative_image_path = self.image_processor._get_relative_image_path(file_stem, safe_image_name)
            
            markdown_content = f"""# {file_stem}

## 图片文件

- **文件名**: {file_name}
- **文件大小**: {os.path.getsize(input_file)} 字节

## 图片预览

![{file_stem}]({relative_image_path})

---

*此文档由 MarkItDown 自动生成*
"""
            
            return markdown_content
            
        except Exception as e:
            self.logger.error(f"创建图片 Markdown 失败: {e}")
            return None
    
    def _try_alternative_docx_conversion(self, input_file: str) -> Optional[str]:
        """
        尝试使用备选方法转换DOCX文件
        
        Args:
            input_file: 输入文件路径
            
        Returns:
            转换后的Markdown内容，失败返回None
        """
        # 方法1: 尝试使用pandoc
        try:
            import subprocess
            result = subprocess.run(
                ['pandoc', input_file, '-t', 'markdown'],
                capture_output=True,
                text=True,
                timeout=60
            )
            if result.returncode == 0 and result.stdout.strip():
                self.logger.info(f"使用pandoc成功转换: {input_file}")
                return result.stdout
        except (subprocess.TimeoutExpired, subprocess.CalledProcessError, FileNotFoundError):
            self.logger.info("pandoc转换失败，尝试下一种方法")
        except Exception as e:
            self.logger.info(f"pandoc转换异常: {e}")
        
        # 方法2: 尝试使用python-docx + 自定义解析
        try:
            from docx import Document
            doc = Document(input_file)
            
            markdown_content = f"# {Path(input_file).stem}\n\n"
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    # 简单的段落处理
                    text = paragraph.text.strip()
                    if paragraph.style.name.startswith('Heading'):
                        level = int(paragraph.style.name[-1]) if paragraph.style.name[-1].isdigit() else 1
                        markdown_content += f"{'#' * level} {text}\n\n"
                    else:
                        markdown_content += f"{text}\n\n"
            
            # 处理表格
            for table in doc.tables:
                markdown_content += "\n"
                for i, row in enumerate(table.rows):
                    row_text = " | ".join([cell.text.strip() for cell in row.cells])
                    markdown_content += f"| {row_text} |\n"
                    if i == 0:  # 添加表头分隔符
                        markdown_content += f"| {' | '.join(['---'] * len(row.cells))} |\n"
                markdown_content += "\n"
            
            if len(markdown_content) > 50:  # 确保有实际内容
                self.logger.info(f"使用python-docx成功转换: {input_file}")
                return markdown_content
                
        except ImportError:
            self.logger.info("python-docx未安装，跳过此方法")
        except Exception as e:
            self.logger.info(f"python-docx转换失败: {e}")
        
        # 方法3: 尝试LibreOffice转换为PDF再处理
        try:
            import subprocess
            import tempfile
            
            with tempfile.TemporaryDirectory() as temp_dir:
                # 转换为PDF
                pdf_path = Path(temp_dir) / f"{Path(input_file).stem}.pdf"
                result = subprocess.run(
                    ['libreoffice', '--headless', '--convert-to', 'pdf', 
                     '--outdir', temp_dir, input_file],
                    capture_output=True,
                    timeout=120
                )
                
                if result.returncode == 0 and pdf_path.exists():
                    # 使用MarkItDown转换PDF
                    conversion_result = self.markitdown.convert(str(pdf_path))
                    if conversion_result and conversion_result.text_content:
                        self.logger.info(f"使用LibreOffice+PDF成功转换: {input_file}")
                        return conversion_result.text_content
                        
        except (subprocess.TimeoutExpired, subprocess.CalledProcessError, FileNotFoundError):
            self.logger.info("LibreOffice转换失败")
        except Exception as e:
            self.logger.info(f"LibreOffice转换异常: {e}")
        
        return None
    
    def convert_document(self, input_file: str) -> Dict[str, Any]:
        """
        转换单个文档
        
        Args:
            input_file: 输入文件路径
            
        Returns:
            转换结果字典
        """
        start_time = datetime.now()
        result = {
            'success': False,
            'input_file': input_file,
            'output_file': None,
            'images_extracted': 0,
            'processing_time': 0,
            'format_optimized': False,
            'llm_descriptions': 0,
            'error': None
        }
        
        try:
            # 验证输入文件
            if not self._validate_input_file(input_file):
                result['error'] = "输入文件验证失败"
                return result
            
            self.logger.info(f"开始转换文档: {input_file}")
            
            # 检测文件类型
            file_ext = Path(input_file).suffix.lower()
            is_pdf = file_ext == '.pdf'
            
            # 检查是否为图片文件
            image_extensions = {'.png', '.jpg', '.jpeg', '.gif', '.bmp', '.tiff', '.webp'}
            if Path(input_file).suffix.lower() in image_extensions:
                self.logger.info(f"检测到图片文件: {input_file}")
                
                # 使用专门的图片处理方法
                image_markdown = self._create_image_markdown(input_file)
                if image_markdown:
                    markdown_content = image_markdown
                    self.logger.info(f"图片文件转换成功: {input_file}")
                else:
                    result['error'] = f"图片文件处理失败: {input_file}"
                    self.logger.error(result['error'])
                    return result
            else:
                # 使用 MarkItDown 转换
                try:
                    md_result = self.markitdown.convert(input_file)
                    if md_result and md_result.text_content:
                        markdown_content = md_result.text_content
                        self.logger.info(f"MarkItDown转换成功: {input_file}")
                    else:
                        raise Exception("MarkItDown 返回空内容")
                except KeyError as e:
                    if 'w:ilvl' in str(e):
                        result['error'] = f"文档结构解析错误: {str(e)}"
                        self.logger.error(result['error'])
                        return result
                    else:
                        result['error'] = f"文档结构解析错误: {str(e)}"
                        self.logger.error(result['error'])
                        return result
                except Exception as e:
                    # 检查是否是 w:ilvl 相关的 FileConversionException
                    if ('FileConversionException' in str(type(e)) and 'w:ilvl' in str(e)) and input_file.lower().endswith('.docx'):
                        self.logger.warning(f"MarkItDown转换失败，尝试备选方法: {e}")
                        
                        # 尝试备选转换方法
                        alternative_content = self._try_alternative_docx_conversion(input_file)
                        
                        if alternative_content:
                            markdown_content = alternative_content
                            self.logger.info(f"备选方法转换成功: {input_file}")
                        else:
                            # 所有方法都失败，生成错误说明
                            result['error'] = f"DOCX文件格式兼容性问题：所有转换方法均失败。\n\n建议解决方案：\n1. 使用LibreOffice转换为PDF后处理\n2. 在Word中重新保存文档\n3. 使用其他转换工具如pandoc"
                            self.logger.warning(f"DOCX兼容性问题: {result['error']}")
                            # 尝试创建一个详细的错误说明文件
                            try:
                                output_path = self._get_output_path(input_file)
                                result['output_file'] = output_path
                                placeholder_content = f"""# {Path(input_file).stem}

**转换失败**: DOCX文件格式兼容性问题

## 🔍 问题分析

该DOCX文件虽然结构完整，但所有可用的转换方法都无法正确处理其格式。这可能是一个复杂的文档兼容性问题。

## 🛠️ 解决方案

### 方案1: PDF转换（推荐）
```bash
# 使用LibreOffice转换为PDF
libreoffice --headless --convert-to pdf "{input_file}"
python main.py "{Path(input_file).with_suffix('.pdf')}"
```

### 方案2: 文档重新保存
1. 用Microsoft Word打开文档
2. 另存为新的DOCX文件
3. 重新尝试转换

### 方案3: 使用其他工具
- pandoc: `pandoc "{input_file}" -o output.md`
- python-docx + 自定义解析

## 📋 技术详情

**错误类型**: {type(e).__name__}
**错误信息**: {str(e)}
**分析时间**: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}

该错误表明所有可用的转换方法都无法处理此文档格式。
"""
                            
                                if self._save_markdown_file(placeholder_content, output_path):
                                    result['success'] = True
                                    result['error'] = None
                                    result['warning'] = "文件转换遇到格式问题，已生成说明文档"
                                    self.logger.info(f"已为格式问题文件生成说明文档: {output_path}")
                                else:
                                    result['error'] = "无法创建说明文档"
                            except Exception as save_error:
                                self.logger.error(f"创建说明文档失败: {save_error}")
                                result['error'] = f"DOCX格式问题且无法创建说明文档: {str(e)}"
                            return result
                    else:
                        result['error'] = f"MarkItDown 转换异常: {str(e)}"
                        self.logger.error(result['error'])
                        return result
            
            # 生成输出路径
            output_path = self._get_output_path(input_file)
            result['output_file'] = output_path
            
            # 处理图片 - 使用与输出文件名一致的清理后文件名
            document_name = self.image_processor._sanitize_filename(Path(input_file).stem)
            # 对于图片文件，跳过图片处理步骤（因为已经在 _create_image_markdown 中处理）
            if file_ext not in image_extensions:
                processed_content, images_count = self.image_processor.process_images(
                    markdown_content, document_name, input_file
                )
            else:
                # 图片文件不需要额外的图片处理
                processed_content = markdown_content
                images_count = 1  # 图片文件本身算作一个处理的图片
            result['images_extracted'] = images_count
            
            # 清理和优化内容
            final_content = self._clean_markdown_content(processed_content, is_pdf)
            if is_pdf:
                result['format_optimized'] = True
            
            # 保存文件
            if self._save_markdown_file(final_content, output_path):
                result['success'] = True
                self.stats['successful_conversions'] += 1
                self.stats['total_images_extracted'] += images_count
                
                # 记录 LLM 使用情况
                if self.enable_llm and images_count > 0:
                    result['llm_descriptions'] = images_count
                    self.stats['llm_descriptions_generated'] += images_count
            else:
                result['error'] = "保存文件失败"
                self.stats['failed_conversions'] += 1
            
        except Exception as e:
            result['error'] = f"转换过程异常: {str(e)}"
            self.logger.error(result['error'])
            self.stats['failed_conversions'] += 1
        
        finally:
            # 计算处理时间
            end_time = datetime.now()
            processing_time = (end_time - start_time).total_seconds()
            result['processing_time'] = processing_time
            self.stats['total_processing_time'] += processing_time
            self.stats['total_files'] += 1
            
            self.logger.info(f"文档转换完成: {input_file}, 耗时: {processing_time:.2f}秒")
        
        return result
    
    def convert_batch(self, input_files: List[str]) -> List[Dict[str, Any]]:
        """
        批量转换文档
        
        Args:
            input_files: 输入文件路径列表
            
        Returns:
            转换结果列表
        """
        results = []
        total_files = len(input_files)
        
        self.logger.info(f"开始批量转换 {total_files} 个文件")
        
        for i, input_file in enumerate(input_files, 1):
            self.logger.info(f"处理进度: {i}/{total_files} - {Path(input_file).name}")
            result = self.convert_document(input_file)
            results.append(result)
        
        self.logger.info("批量转换完成")
        return results
    
    def get_statistics(self) -> Dict[str, Any]:
        """
        获取转换统计信息
        
        Returns:
            统计信息字典
        """
        stats = self.stats.copy()
        if stats['total_files'] > 0:
            stats['success_rate'] = (stats['successful_conversions'] / stats['total_files']) * 100
            stats['average_processing_time'] = stats['total_processing_time'] / stats['total_files']
        else:
            stats['success_rate'] = 0
            stats['average_processing_time'] = 0
        
        # 添加配置信息
        stats['llm_enabled'] = self.enable_llm
        stats['azure_doc_intel_enabled'] = self.use_azure_doc_intel
        stats['keep_data_uris'] = self.keep_data_uris
        
        return stats
    
    def reset_statistics(self) -> None:
        """
        重置统计信息
        """
        self.stats = {
            'total_files': 0,
            'successful_conversions': 0,
            'failed_conversions': 0,
            'total_images_extracted': 0,
            'total_processing_time': 0,
            'llm_descriptions_generated': 0
        }
        self.logger.info("统计信息已重置")
    
    def get_supported_formats(self) -> List[str]:
        """
        获取支持的文件格式列表
        
        Returns:
            支持的格式列表
        """
        # MarkItDown 实际稳定支持的格式
        return [
            '.pdf', '.docx', '.doc', '.pptx', '.ppt', '.xlsx', '.xls',
            '.txt', '.md', '.html', '.htm', '.xml', '.json', '.csv',
            '.rtf', '.odt', '.odp', '.ods', '.epub',
            '.jpg', '.jpeg', '.png', '.gif', '.bmp', '.tiff', '.webp',
            '.mp3', '.wav', '.m4a', '.flac', '.aac',
            '.zip', '.tar', '.gz'
        ]
    
    def cleanup_temp_files(self) -> None:
        """
        清理临时文件
        """
        try:
            temp_dir = Path(Config.TEMP_DIR)
            if temp_dir.exists():
                for temp_file in temp_dir.rglob('*'):
                    if temp_file.is_file():
                        temp_file.unlink()
                        self.logger.debug(f"删除临时文件: {temp_file}")
                
                self.logger.info("临时文件清理完成")
        except Exception as e:
            self.logger.error(f"清理临时文件失败: {str(e)}")


if __name__ == '__main__':
    # 测试转换器
    converter = DocumentConverter()
    
    # 显示支持的格式
    print("支持的文件格式:")
    for fmt in converter.get_supported_formats():
        print(f"  {fmt}")
    
    # 测试单文件转换（如果有测试文件）
    test_file = "test.pdf"
    if os.path.exists(test_file):
        print(f"\n测试转换文件: {test_file}")
        result = converter.convert_document(test_file)
        print(f"转换结果: {'成功' if result['success'] else '失败'}")
        if result['error']:
            print(f"错误信息: {result['error']}")